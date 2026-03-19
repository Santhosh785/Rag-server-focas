import re
import io
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_shading(cell, color):
    """Sets background color for a table cell."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_word_table(container, table_lines):
    """Parses ASCII pipes into a proportional Word table."""
    data_rows = []
    for line in table_lines:
        line_s = line.strip()
        if not line_s.startswith('|'):
            continue
        cells = [c.strip() for c in line_s.split('|')]
        # Remove first empty before first pipe
        if len(cells) > 0 and cells[0] == "": cells = cells[1:]
        # Remove last empty after last pipe
        if len(cells) > 0 and cells[-1] == "": cells = cells[:-1]
        
        # Don't add markdown dividers consisting of just dashes e.g. |---|---|
        if all(re.match(r'^[-:\s]*$', c) for c in cells):
            continue

        data_rows.append(cells)
    
    if not data_rows: return
    # Remove rows that are entirely empty
    data_rows = [r for r in data_rows if any(c != "" for c in r)]
    if not data_rows: return

    # Ensure consistent number of columns
    num_cols = max(len(row) for row in data_rows)
    for r in data_rows:
        while len(r) < num_cols:
            r.append("")

    # Filter out empty columns
    cols_to_keep = []
    for j in range(num_cols):
        if any(row[j] != "" for row in data_rows):
            cols_to_keep.append(j)

    if not cols_to_keep:
        return

    filtered_data = []
    for row in data_rows:
        filtered_data.append([row[j] for j in cols_to_keep])

    data_rows = filtered_data
    num_cols = len(cols_to_keep)

    table = container.add_table(rows=len(data_rows), cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    
    for i, row in enumerate(data_rows):
        # Prevent row from breaking across pages
        tr = table.rows[i]._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), 'true')
        trPr.append(cantSplit)
        
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = val
            for p in cell.paragraphs:
                p.style.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)

    # Add a small buffer paragraph after table
    try:
        p_buf = container.add_paragraph()
        p_buf.paragraph_format.space_after = Pt(4)
    except:
        pass

def add_formatted_content(container, text, left_indent=None):
    """Detects ASCII tables in text and converts them to native Word tables."""
    lines = text.split('\n')
    table_lines = []
    
    for line in lines:
        stripped = line.strip()
        is_table_row = (stripped.startswith('+') and stripped.endswith('+')) or \
                      (stripped.startswith('|') and stripped.endswith('|'))
        
        if is_table_row:
            table_lines.append(line)
        else:
            if table_lines:
                create_word_table(container, table_lines)
                table_lines = []
            if stripped: 
                p = container.add_paragraph(stripped)
                p.style.font.size = Pt(10)
                if left_indent: p.paragraph_format.left_indent = left_indent
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.0

    if table_lines:
        create_word_table(container, table_lines)
