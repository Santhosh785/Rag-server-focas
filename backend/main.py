import os
import io
import re
import pandas as pd
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
from fastapi.middleware.cors import CORSMiddleware
from pymongo import MongoClient
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import logging
from datetime import datetime, timezone
from dotenv import load_dotenv

load_dotenv()

# --- Config ---
MONGODB_URI = os.environ.get("MONGODB_URI")
if not MONGODB_URI:
    raise RuntimeError("MONGODB_URI not found in environment")

DB_NAME = "exam_db"
COLLECTION = "questions"

# --- Setup ---
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- DB Connection ---
client = MongoClient(MONGODB_URI)
db = client[DB_NAME]
col = db[COLLECTION]

# --- Helper for Styling ---
def set_cell_shading(cell, color):
    """Sets background color for a table cell."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def clean_val(val):
    if pd.isna(val): return ""
    if isinstance(val, (float, int)):
        if float(val).is_integer(): return str(int(val))
        return str(val)
    return str(val).strip()


import zipfile

def build_paper_bundle_from_df(df: pd.DataFrame) -> io.BytesIO:
    # 1. Initialize Word Docs
    qp_doc = Document()
    ak_doc = Document()
    
    # Common margins
    for doc in [qp_doc, ak_doc]:
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.65)
            section.right_margin = Inches(0.65)

    # Pull Metadata
    first_row = df.dropna(subset=['level', 'subject']).iloc[0] if not df.empty else None
    paper_level = clean_val(first_row.get('level', 'CA INTERMEDIATE')) if first_row is not None else 'CA INTERMEDIATE'
    paper_level = paper_level.upper()
    paper_subject = clean_val(first_row.get('subject', 'SUBJECT')) if first_row is not None else 'SUBJECT'
    paper_subject = paper_subject.upper()

    # Total Marks
    total_marks = 100
    if 'marks' in df.columns:
        m_numeric = pd.to_numeric(df['marks'], errors='coerce').fillna(0)
        total_marks = int(m_numeric.sum())
        if total_marks <= 0: total_marks = 100

    # --- Setup Headers for Both ---
    for doc, title_tag in [(qp_doc, "MODEL PAPER"), (ak_doc, "ANSWER KEY")]:
        # TOP HEADER BAR
        bar_table = doc.add_table(rows=1, cols=1)
        bar_table.width = Inches(7.3)
        cell = bar_table.rows[0].cells[0]
        set_cell_shading(cell, "002060") # Dark Blue
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("FOCAS EDU — LAST ATTEMPT KIT PRO")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255) # White
        run.font.size = Pt(14)

        # EXAM TITLES
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        p_exam = doc.add_paragraph(f"{paper_level} EXAMINATION")
        p_exam.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_exam.add_run().bold = True
        p_exam.runs[0].font.size = Pt(13)
        
        p_sub = doc.add_paragraph(f"PAPER: {paper_subject}\nFULL TEST — {title_tag}")
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.runs[0].bold = True
        p_sub.runs[0].font.size = Pt(11)

        # STATS ROW
        stats_table = doc.add_table(rows=1, cols=3)
        stats_table.width = Inches(7.3)
        stats_table.style = 'Table Grid'
        stats_data = [f"Total Marks: {total_marks}", "Time: 3 Hours", "Date: _________"]
        for i, text in enumerate(stats_data):
            c = stats_table.rows[0].cells[i]
            set_cell_shading(c, "CFE2F3")
            c.text = text
            cp = c.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].bold = True
            cp.runs[0].font.size = Pt(10)

    # --- INSTRUCTIONS (Only for QP) ---
    qp_doc.add_paragraph().paragraph_format.space_after = Pt(4)
    inst_bar = qp_doc.add_table(rows=1, cols=1)
    inst_bar.width = Inches(7.3)
    inst_cell = inst_bar.rows[0].cells[0]
    set_cell_shading(inst_cell, "FFF2CC") # Pale Yellow
    ip = inst_cell.paragraphs[0]
    ip.add_run("GENERAL INSTRUCTIONS:").bold = True
    for inst in ["1. All questions are COMPULSORY.", "2. Marks are indicated against each question in brackets [ ].", "3. Answers should be based on relevant Study Material and standards."]:
        p_inst = qp_doc.add_paragraph(inst)
        p_inst.style.font.size = Pt(9)
        p_inst.paragraph_format.left_indent = Inches(0.2)
        p_inst.paragraph_format.space_after = Pt(0)

    # --- PART BAR ---
    for doc in [qp_doc, ak_doc]:
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        part_table = doc.add_table(rows=1, cols=1)
        part_table.width = Inches(7.3)
        part_cell = part_table.rows[0].cells[0]
        set_cell_shading(part_cell, "002060") # Dark Blue
        pp = part_cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_p = pp.add_run("PART — I")
        run_p.bold = True
        run_p.font.color.rgb = RGBColor(255, 255, 255)
        run_p.font.size = Pt(11)

    # 3. Retrieve and Format Questions & Answers
    questions_found = 0
    for index, row in df.iterrows():
        q_num = clean_val(row.get('question_number'))
        subject = clean_val(row.get('subject'))
        level = clean_val(row.get('level'))
        chapter = clean_val(row.get('chapter_number'))
        unit = clean_val(row.get('unit'))
        marks = clean_val(row.get('marks'))

        if not q_num: continue

        query = {
            "level": {"$regex": f"^{re.escape(level)}$", "$options": "i"},
            "subject": {"$regex": f"^{re.escape(subject)}$", "$options": "i"},
            "chapter": {"$regex": f"^{re.escape(chapter)}$", "$options": "i"},
            "question_no": {"$regex": f"^{re.escape(q_num)}$", "$options": "i"}
        }
        if unit: query["unit"] = {"$regex": f"^{re.escape(unit)}$", "$options": "i"}
            
        question_data = col.find_one(query)
        if question_data:
            questions_found += 1
            
            # --- PROCESS QUESTION (for QP) ---
            q_text = question_data.get("question_text", "")
            q_text_clean = re.sub(r'^\s*(?:QUESTION|Question)\s+(?:NO\.?\s+)?\d+[^\n]*\n', '', q_text, flags=re.IGNORECASE).strip()
            # 2. Exhaustively remove all leading metadata blocks (e.g. (MTP...), [RTP...], (PYP...))
            while True:
                # Matches anything starting with ( or [ and ending with ) or ] at the beginning
                meta_match = re.match(r'^([\[\(].*?[\]\)])\s*', q_text_clean, flags=re.IGNORECASE)
                if meta_match:
                    content = meta_match.group(1).upper()
                    # If it contains typical metadata keywords, remove it
                    keywords = ["MTP", "RTP", "PYP", "MARK", "MAY", "NOV", "OCT", "APR", "20"]
                    if any(k in content for k in keywords):
                        q_text_clean = q_text_clean[meta_match.end():].strip()
                        continue
                break

            
            lines = q_text_clean.split('\n')
            first_text_line = ""
            remaining_lines = []
            has_first = False
            for line in lines:
                s = line.strip()
                if not has_first:
                    if not s: continue
                    if s.startswith('+') or s.startswith('|'):
                        remaining_lines.append(line)
                        has_first = True
                    else:
                        first_text_line = s + " "
                        has_first = True
                else: remaining_lines.append(line)

            head_table = qp_doc.add_table(rows=1, cols=2)
            head_table.autofit = False
            head_table.columns[0].width = Inches(5.8)
            head_table.columns[1].width = Inches(1.2)
            c_left = head_table.rows[0].cells[0]
            p_head = c_left.paragraphs[0]
            p_head.paragraph_format.space_before = Pt(16) if questions_found > 1 else Pt(4)
            run_q_label = p_head.add_run(f"Q{questions_found}. ")
            run_q_label.bold = True
            run_q_label.font.size = Pt(11)
            # Note: source_meta is extracted but not added to p_head as per user request to hide (RTP/MTP) metadata

            if first_text_line:
                body_run = p_head.add_run("  " + first_text_line.strip())
                body_run.font.size = Pt(10)

            c_right = head_table.rows[0].cells[1]
            p_marks = c_right.paragraphs[0]
            p_marks.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if marks:
                 m_val = str(marks).strip()
                 if '[' not in m_val and 'Mark' not in m_val: m_val = f"[{m_val} Marks]"
                 mk_run = p_marks.add_run(m_val)
                 mk_run.bold = True
                 mk_run.font.color.rgb = RGBColor(192, 0, 0)
                 mk_run.font.size = Pt(11)

            p_head.paragraph_format.space_after = Pt(2)
            if remaining_lines: add_formatted_content(qp_doc, "\n".join(remaining_lines))

            # --- PROCESS ANSWER (for AK) ---
            ans_text = question_data.get("answer_text", "") or "No answer found in database."
            # Clean leading "Answer" keyword if it matches the pattern
            ans_text_clean = re.sub(r'^\s*Answer\s*\n*', '', ans_text, flags=re.IGNORECASE).strip()
            
            ak_head_p = ak_doc.add_paragraph()
            ak_head_p.paragraph_format.space_before = Pt(16) if questions_found > 1 else Pt(4)
            ak_head_run = ak_head_p.add_run(f"Answer to Q{questions_found}")
            ak_head_run.bold = True
            ak_head_run.font.size = Pt(11)
            ak_head_run.underline = True
            ak_head_p.paragraph_format.space_after = Pt(6)
            
            add_formatted_content(ak_doc, ans_text_clean)
            
        else:
            logger.warning(f"❌ Not found in DB: Q{q_num} (L={level}, S={subject}, Ch={chapter})")

    if questions_found == 0:
        raise HTTPException(status_code=404, detail="No matching questions found in database.")

    # 4. Save both to separate BytesIO
    qp_output = io.BytesIO()
    qp_doc.save(qp_output)
    qp_output.seek(0)
    
    ak_output = io.BytesIO()
    ak_doc.save(ak_output)
    ak_output.seek(0)
    
    # 5. Create ZIP package
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.writestr("01_Question_Paper.docx", qp_output.getvalue())
        zip_file.writestr("02_Answer_Key.docx", ak_output.getvalue())
    
    zip_buffer.seek(0)
    return zip_buffer

@app.post("/api/generate-paper")
async def generate_paper(file: UploadFile = File(...)):
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Please upload an Excel file.")

    try:
        contents = await file.read()
        df_raw = pd.read_excel(io.BytesIO(contents), header=None)
        
        header_row_index = 0
        required_keywords = {"question", "subject", "level", "chapter"}
        for i, row in df_raw.iterrows():
            row_values = [str(val).lower() for val in row.values if pd.notna(val)]
            if sum(1 for kw in required_keywords if any(kw in val for val in row_values)) >= 3:
                header_row_index = i
                break
        
        df = pd.read_excel(io.BytesIO(contents), header=header_row_index)
        df.columns = [str(c).strip().lower() for c in df.columns]
        
        column_mapping = {
            "question_number": ["question", "q_no", "q no", "number"],
            "subject": ["subject", "sub"],
            "level": ["level", "lvl"],
            "chapter_number": ["chapter", "ch_no", "ch no"],
            "unit": ["unit", "unit_no", "u_no", "u no"],
            "marks": ["marks", "mark", "pts", "points"]
        }
        for target, aliases in column_mapping.items():
            if target not in df.columns:
                for col_name in df.columns:
                    if any(alias in col_name for alias in aliases):
                        df.rename(columns={col_name: target}, inplace=True)
                        break
 

        zip_output = build_paper_bundle_from_df(df)
        filename = f"CA_Exam_Package_{datetime.now().strftime('%Y%m%d')}.zip"
        return StreamingResponse(
            zip_output,
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating paper: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class QuestionItem(BaseModel):
    level: str
    subject: str
    chapter_number: str
    unit: Optional[str] = ""
    question_number: str
    marks: Optional[str] = ""

class PaperRequest(BaseModel):
    questions: List[QuestionItem]

class RandomPaperRequest(BaseModel):
    level: str
    subject: str
    chapter_number: Optional[str] = ""
    total_marks: int = 50

@app.post("/api/generate-paper-json")
async def generate_paper_json(data: PaperRequest):
    try:
        df = pd.DataFrame([item.dict() for item in data.questions])
        if df.empty:
            raise HTTPException(status_code=400, detail="No questions provided in list.")
        
        zip_output = build_paper_bundle_from_df(df)
        filename = f"CA_Exam_Package_{datetime.now().strftime('%Y%m%d')}.zip"
        return StreamingResponse(
            zip_output,
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating paper from JSON: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/generate-random-paper")
async def generate_random_paper(data: RandomPaperRequest):
    try:
        query = {
            "level": {"$regex": f"^{re.escape(data.level)}$", "$options": "i"},
            "subject": {"$regex": f"^{re.escape(data.subject)}$", "$options": "i"}
        }
        if data.chapter_number:
            query["chapter"] = {"$regex": f"^{re.escape(data.chapter_number)}$", "$options": "i"}

        # Fetch a large sample
        pipeline = [{"$match": query}, {"$sample": {"size": 200}}]
        results = list(col.aggregate(pipeline))

        if not results:
            raise HTTPException(status_code=404, detail="No matching questions found in DB for this criteria.")

        selected_questions = []
        current_marks = 0

        for q in results:
            if current_marks >= data.total_marks:
                break
            
            q_text = q.get("question_text", "")
            # Extract marks looking for "[... X Marks ...]"
            marks_match = re.search(r'\[.*?(\d+)\s*Marks?', q_text, re.IGNORECASE)
            q_marks = int(marks_match.group(1)) if marks_match and marks_match.group(1).isdigit() else 5

            if current_marks + q_marks > data.total_marks + 2 and current_marks > 0:
                continue
                
            selected_questions.append({
                "level": q.get("level", data.level),
                "subject": q.get("subject", data.subject),
                "chapter_number": q.get("chapter", data.chapter_number or ""),
                "unit": q.get("unit", ""),
                "question_number": q.get("question_no", ""),
                "marks": str(q_marks)
            })
            current_marks += q_marks

        if not selected_questions:
             raise HTTPException(status_code=400, detail="Could not create paper with requested marks.")

        df = pd.DataFrame(selected_questions)
        zip_output = build_paper_bundle_from_df(df)
        filename = f"FOCAS_Random_Package_{datetime.now().strftime('%Y%m%d')}.zip"
        return StreamingResponse(
            zip_output,
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating random paper: {e}")
        raise HTTPException(status_code=500, detail=str(e))

def add_formatted_content(doc, text):
    """Detects ASCII tables in text and converts them to native Word tables."""
    lines = text.split('\n')
    table_lines = []
    
    for line in lines:
        stripped = line.strip()
        is_table_row = (stripped.startswith('+') and stripped.endswith('+')) or \
                      (stripped.startswith('|') and stripped.endswith('|'))
        
        if is_table_row:
            table_lines.append(line)
        else:
            if table_lines:
                create_word_table(doc, table_lines)
                table_lines = []
            if stripped:  # REMOVED `or line == ""` TO COMPACT THE SPACING!
                p = doc.add_paragraph(stripped)
                p.style.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.0

    if table_lines:
        create_word_table(doc, table_lines)

def create_word_table(doc, table_lines):
    """Parses ASCII pipes into a proportional Word table."""
    data_rows = []
    for line in table_lines:
        line_s = line.strip()
        if not line_s.startswith('|'):
            continue
        cells = [c.strip() for c in line_s.split('|')]
        # Remove first empty before first pipe
        if len(cells) > 0 and cells[0] == "": cells = cells[1:]
        # Remove last empty after last pipe
        if len(cells) > 0 and cells[-1] == "": cells = cells[:-1]
        
        # Don't add markdown dividers consisting of just dashes e.g. |---|---|
        if all(re.match(r'^[-:\s]*$', c) for c in cells):
            continue

        data_rows.append(cells)
    
    if not data_rows: return
    # Remove rows that are entirely empty
    data_rows = [r for r in data_rows if any(c != "" for c in r)]
    if not data_rows: return

    # Remove completely empty columns
    num_cols = max(len(row) for row in data_rows)
    for r in data_rows:
        while len(r) < num_cols:
            r.append("")

    cols_to_keep = []
    for j in range(num_cols):
        if any(row[j] != "" for row in data_rows):
            cols_to_keep.append(j)

    if not cols_to_keep:
        return

    filtered_data = []
    for row in data_rows:
        filtered_data.append([row[j] for j in cols_to_keep])

    data_rows = filtered_data
    num_cols = len(cols_to_keep)

    table = doc.add_table(rows=len(data_rows), cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for i, row in enumerate(data_rows):
        # Prevent row from breaking across pages
        tr = table.rows[i]._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), 'true')
        trPr.append(cantSplit)
        
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = val
            for p in cell.paragraphs:
                p.style.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)

    # Add a small buffer paragraph after table
    try:
        p_buf = doc.add_paragraph()
        p_buf.paragraph_format.space_after = Pt(4)
    except:
        pass

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
