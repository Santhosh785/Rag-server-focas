import io
import re
import zipfile
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import logging

from backend.utils.docx_utils import (
    set_cell_shading,
    add_formatted_content
)
from backend.utils.text_utils import (
    clean_val,
    clean_question_text,
    arabic_to_roman
)

logger = logging.getLogger(__name__)

def build_paper_bundle_from_df(df: pd.DataFrame, col) -> io.BytesIO:
    # 1. Initialize Word Docs
    qp_doc = Document()
    ak_doc = Document()
    
    # ── Style Configuration ──
    for doc in [qp_doc, ak_doc]:
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Cambria'
        font.size = Pt(11)
        
        # Common margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
            
            # Add Footer (Page X of Y)
            footer = section.footer
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fp.add_run("Page ")
            # Note: docx doesn't easily support "Page X of Y" fields without complex OXML
            # We'll just leave Page placeholder or current page number
            fp.style.font.size = Pt(9)

    # Pull Metadata
    first_row = df.dropna(subset=['level', 'subject']).iloc[0] if not df.empty else None
    paper_level = clean_val(first_row.get('level', 'CA INTERMEDIATE')) if first_row is not None else 'CA INTERMEDIATE'
    paper_level = paper_level.upper()
    paper_subject = clean_val(first_row.get('subject', 'SUBJECT')) if first_row is not None else 'SUBJECT'
    paper_subject = paper_subject.upper()

    # Total Marks
    total_marks = 100
    if 'marks' in df.columns:
        m_numeric = pd.to_numeric(df['marks'], errors='coerce').fillna(0)
        total_marks = int(m_numeric.sum())
        if total_marks <= 0: total_marks = 100

    # --- Setup Headers for Both ---
    for doc, title_tag in [(qp_doc, "MODEL PAPER"), (ak_doc, "SUGGESTED ANSWERS")]:
        # TOP HEADER BAR
        bar_table = doc.add_table(rows=1, cols=1)
        bar_table.width = Inches(7.1)
        bar_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = bar_table.rows[0].cells[0]
        set_cell_shading(cell, "002060") # Dark Blue
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("FOCAS EDU — LAST ATTEMPT KIT PRO")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255) # White
        run.font.size = Pt(14)

        # EXAM TITLES
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        p_exam = doc.add_paragraph()
        p_exam.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_e = p_exam.add_run(f"{paper_level} EXAMINATION")
        run_e.bold = True
        run_e.font.size = Pt(12)
        
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_s = p_sub.add_run(f"PAPER: {paper_subject}\n({title_tag})")
        run_s.bold = True
        run_s.font.size = Pt(11)

        # STATS ROW
        stats_table = doc.add_table(rows=1, cols=3)
        stats_table.width = Inches(7.1)
        stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        stats_table.style = 'Table Grid'
        stats_data = [f"Total Marks: {total_marks}", "Time: 3 Hours", "Date: _________"]
        for i, text in enumerate(stats_data):
            c = stats_table.rows[0].cells[i]
            set_cell_shading(c, "F2F2F2")
            c.text = text
            cp = c.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cp.runs[0]
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Cambria'

    # --- INSTRUCTIONS (Only for QP) ---
    qp_doc.add_paragraph().paragraph_format.space_after = Pt(8)
    inst_p = qp_doc.add_paragraph()
    run_inst = inst_p.add_run("GENERAL INSTRUCTIONS:")
    run_inst.bold = True
    run_inst.underline = True
    
    for idx, inst in enumerate(["All questions are COMPULSORY.", "Marks are indicated against each question in brackets [ ].", "Answers should be based on relevant Study Material and standards."], 1):
        p_inst = qp_doc.add_paragraph(f"{idx}. {inst}")
        p_inst.style.font.size = Pt(9.5)
        p_inst.paragraph_format.left_indent = Inches(0.2)
        p_inst.paragraph_format.space_after = Pt(2)

    # --- PART BAR ---
    for doc in [qp_doc, ak_doc]:
        doc.add_paragraph().paragraph_format.space_after = Pt(10)
        part_table = doc.add_table(rows=1, cols=1)
        part_table.width = Inches(7.1)
        part_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        part_cell = part_table.rows[0].cells[0]
        set_cell_shading(part_cell, "002060") # Dark Blue
        pp = part_cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_p = pp.add_run("PART — I")
        run_p.bold = True
        run_p.font.color.rgb = RGBColor(255, 255, 255)
        run_p.font.size = Pt(11)

    # 3. Retrieve and Format Questions & Answers
    questions_found = 0
    for index, row in df.iterrows():
        q_num = clean_val(row.get('question_number'))
        subject = clean_val(row.get('subject'))
        level = clean_val(row.get('level'))
        chapter = clean_val(row.get('chapter_number'))
        unit = clean_val(row.get('unit'))
        marks = clean_val(row.get('marks'))

        if not q_num: continue

        # Base query
        base_query = {
            "level": {"$regex": f"^{re.escape(level)}$", "$options": "i"},
            "subject": {"$regex": f"^{re.escape(subject)}$", "$options": "i"},
            "chapter": {"$regex": f"^{re.escape(chapter)}$", "$options": "i"},
            "question_no": {"$regex": f"^{re.escape(q_num)}$", "$options": "i"}
        }

        if unit:
            # Try to convert Arabic to Roman (e.g., "2" -> "II")
            unit_roman = arabic_to_roman(unit)
            # Search for both Arabic and Roman versions using $or
            if unit_roman and unit_roman != unit:
                query = {
                    "$and": [
                        base_query,
                        {
                            "$or": [
                                {"unit": {"$regex": f"^{re.escape(unit)}$", "$options": "i"}},
                                {"unit": {"$regex": f"^{re.escape(unit_roman)}$", "$options": "i"}}
                            ]
                        }
                    ]
                }
            else:
                query = base_query
                query["unit"] = {"$regex": f"^{re.escape(unit)}$", "$options": "i"}
        else:
            query = base_query
            
        question_data = col.find_one(query)
        if question_data:
            questions_found += 1
            
            # --- PROCESS QUESTION (for QP) ---
            q_text = question_data.get("question_text", "")
            q_text_clean = clean_question_text(q_text)
            
            lines = q_text_clean.split('\n')
            first_text_line = ""
            remaining_lines = []
            has_first = False
            for line in lines:
                s = line.strip()
                if not has_first:
                    if not s: continue
                    if s.startswith('+') or s.startswith('|'):
                        remaining_lines.append(line)
                        has_first = True
                    else:
                        first_text_line = s
                        has_first = True
                else: remaining_lines.append(line)

            head_table = qp_doc.add_table(rows=1, cols=2)
            head_table.autofit = False
            head_table.width = Inches(7.1)
            head_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            c_left = head_table.rows[0].cells[0]
            c_right = head_table.rows[0].cells[1]
            c_left.width = Inches(6.1)
            c_right.width = Inches(1.0)
            
            # Setup first paragraph in the cell
            p_head = c_left.paragraphs[0]
            # More space between questions
            p_head.paragraph_format.space_before = Pt(18) if questions_found > 1 else Pt(10)
            p_head.paragraph_format.line_spacing = 1.15
            
            # Label "Q1. " with hanging indent
            p_head.paragraph_format.left_indent = Inches(0.45)
            p_head.paragraph_format.first_line_indent = Inches(-0.45)
            
            run_q_label = p_head.add_run(f"Q{questions_found}. ")
            run_q_label.bold = True
            run_q_label.font.size = Pt(11)
            
            if first_text_line:
                body_run = p_head.add_run(first_text_line.strip())
                body_run.font.size = Pt(10.5)

            # Keep remaining lines INSIDE the same cell for consistent alignment
            if remaining_lines:
                add_formatted_content(c_left, "\n".join(remaining_lines), left_indent=Inches(0.45))

            # Marks in the right cell
            p_marks = c_right.paragraphs[0]
            p_marks.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_marks.paragraph_format.space_before = Pt(18) if questions_found > 1 else Pt(10)
            if marks:
                 m_val = str(marks).strip()
                 if '[' not in m_val and 'Mark' not in m_val: m_val = f"[{m_val} Marks]"
                 mk_run = p_marks.add_run(m_val)
                 mk_run.bold = True
                 # Deep red/maroon for professional look
                 mk_run.font.color.rgb = RGBColor(128, 0, 0)
                 mk_run.font.size = Pt(10.5)
                 mk_run.font.italic = True

            p_head.paragraph_format.space_after = Pt(4)

            # --- PROCESS ANSWER (for AK) ---
            ans_text = question_data.get("answer_text", "") or "No answer found in database."
            # Clean leading "Answer" keyword if it matches the pattern
            ans_text_clean = re.sub(r'^\s*Answer\s*\n*', '', ans_text, flags=re.IGNORECASE).strip()
            
            ak_head_p = ak_doc.add_paragraph()
            ak_head_p.paragraph_format.space_before = Pt(20) if questions_found > 1 else Pt(10)
            
            ak_head_run = ak_head_p.add_run(f"Answer to Q{questions_found}")
            ak_head_run.bold = True
            ak_head_run.font.size = Pt(11.5)
            ak_head_run.font.color.rgb = RGBColor(0, 32, 96) # Dark Blue
            ak_head_p.paragraph_format.space_after = Pt(8)
            
            add_formatted_content(ak_doc, ans_text_clean)
            
            # Divider between answers in AK
            ak_doc.add_paragraph("─" * 40).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        else:
            logger.warning(f"❌ Not found in DB: Q{q_num} (L={level}, S={subject}, Ch={chapter})")

    # Final "End of Paper"
    for doc in [qp_doc, ak_doc]:
        end_p = doc.add_paragraph()
        end_p.paragraph_format.space_before = Pt(30)
        end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_end = end_p.add_run("--- END OF PAPER ---")
        run_end.bold = True
        run_end.font.size = Pt(10)

    if questions_found == 0:
        return None

    # 4. Save both to separate BytesIO
    qp_output = io.BytesIO()
    qp_doc.save(qp_output)
    qp_output.seek(0)
    
    ak_output = io.BytesIO()
    ak_doc.save(ak_output)
    ak_output.seek(0)
    
    # 5. Create ZIP package
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.writestr("01_Question_Paper.docx", qp_output.getvalue())
        zip_file.writestr("02_Answer_Key.docx", ak_output.getvalue())
    
    zip_buffer.seek(0)
    return zip_buffer
